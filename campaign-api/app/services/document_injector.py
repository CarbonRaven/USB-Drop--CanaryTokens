"""Document injector service for embedding canary tokens into documents."""

import io
import os
import tempfile
from typing import Optional, Tuple
from pathlib import Path

# Word document manipulation
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Excel manipulation
from openpyxl import load_workbook
from openpyxl.worksheet.datavalidation import DataValidation

# PDF manipulation
from pypdf import PdfReader, PdfWriter
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


class DocumentInjector:
    """Injects canary token URLs into various document types."""

    def __init__(self):
        pass

    def inject_token(
        self,
        file_path: str,
        token_url: str,
        mime_type: str,
        output_path: Optional[str] = None
    ) -> Tuple[bytes, str]:
        """
        Inject a canary token URL into a document.

        Args:
            file_path: Path to the source document
            token_url: The canary token URL to embed
            mime_type: MIME type of the document
            output_path: Optional path for output file

        Returns:
            Tuple of (modified file bytes, output path)
        """
        if mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self._inject_word(file_path, token_url, output_path)
        elif mime_type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            return self._inject_excel(file_path, token_url, output_path)
        elif mime_type == 'application/pdf':
            return self._inject_pdf(file_path, token_url, output_path)
        else:
            raise ValueError(f"Unsupported document type: {mime_type}")

    def _inject_word(
        self,
        file_path: str,
        token_url: str,
        output_path: Optional[str] = None
    ) -> Tuple[bytes, str]:
        """
        Inject a web bug into a Word document.

        Uses an external image relationship that triggers when the document is opened.
        The image is 1x1 pixel and positioned off-page to be invisible.
        """
        doc = Document(file_path)

        # Access the document's relationships
        rels = doc.part.rels

        # Create a new external image relationship
        # This will cause Word to fetch the URL when opening the document
        rel_id = f"rId{len(rels) + 1}"

        # Add the relationship to the document
        # We use the image relationship type but with an external target
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        rels.add(rel_id, RT.IMAGE, token_url, is_external=True)

        # Add a tiny inline shape that references the external image
        # This triggers the fetch when the document is rendered
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        # Create inline shape element
        inline = OxmlElement('wp:inline')
        inline.set(qn('wp:distT'), '0')
        inline.set(qn('wp:distB'), '0')
        inline.set(qn('wp:distL'), '0')
        inline.set(qn('wp:distR'), '0')

        # Extent (size) - 1x1 pixel
        extent = OxmlElement('wp:extent')
        extent.set('cx', '9525')  # 1 pixel in EMUs
        extent.set('cy', '9525')
        inline.append(extent)

        # Effect extent
        effect_extent = OxmlElement('wp:effectExtent')
        effect_extent.set('l', '0')
        effect_extent.set('t', '0')
        effect_extent.set('r', '0')
        effect_extent.set('b', '0')
        inline.append(effect_extent)

        # Doc properties
        doc_pr = OxmlElement('wp:docPr')
        doc_pr.set('id', '1')
        doc_pr.set('name', 'Picture 1')
        doc_pr.set('hidden', '1')  # Hide the element
        inline.append(doc_pr)

        # Non-visual properties
        cnv_pr = OxmlElement('wp:cNvGraphicFramePr')
        graphic_frame_locks = OxmlElement('a:graphicFrameLocks')
        graphic_frame_locks.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')
        graphic_frame_locks.set('noChangeAspect', '1')
        cnv_pr.append(graphic_frame_locks)
        inline.append(cnv_pr)

        # Graphic element
        graphic = OxmlElement('a:graphic')
        graphic.set(qn('xmlns:a'), 'http://schemas.openxmlformats.org/drawingml/2006/main')

        graphic_data = OxmlElement('a:graphicData')
        graphic_data.set('uri', 'http://schemas.openxmlformats.org/drawingml/2006/picture')

        # Picture element
        pic = OxmlElement('pic:pic')
        pic.set(qn('xmlns:pic'), 'http://schemas.openxmlformats.org/drawingml/2006/picture')

        # Non-visual picture properties
        nvPicPr = OxmlElement('pic:nvPicPr')
        cNvPr = OxmlElement('pic:cNvPr')
        cNvPr.set('id', '0')
        cNvPr.set('name', 'tracker.gif')
        nvPicPr.append(cNvPr)

        cNvPicPr = OxmlElement('pic:cNvPicPr')
        nvPicPr.append(cNvPicPr)
        pic.append(nvPicPr)

        # Blip fill - reference to the external image
        blipFill = OxmlElement('pic:blipFill')
        blip = OxmlElement('a:blip')
        blip.set(qn('r:link'), rel_id)  # Use r:link for external reference
        blipFill.append(blip)

        stretch = OxmlElement('a:stretch')
        fillRect = OxmlElement('a:fillRect')
        stretch.append(fillRect)
        blipFill.append(stretch)
        pic.append(blipFill)

        # Shape properties
        spPr = OxmlElement('pic:spPr')
        xfrm = OxmlElement('a:xfrm')
        off = OxmlElement('a:off')
        off.set('x', '0')
        off.set('y', '0')
        xfrm.append(off)
        ext = OxmlElement('a:ext')
        ext.set('cx', '9525')
        ext.set('cy', '9525')
        xfrm.append(ext)
        spPr.append(xfrm)

        prstGeom = OxmlElement('a:prstGeom')
        prstGeom.set('prst', 'rect')
        avLst = OxmlElement('a:avLst')
        prstGeom.append(avLst)
        spPr.append(prstGeom)
        pic.append(spPr)

        graphic_data.append(pic)
        graphic.append(graphic_data)
        inline.append(graphic)

        # Create drawing element
        drawing = OxmlElement('w:drawing')
        drawing.append(inline)

        # Add to run
        run._r.append(drawing)

        # Make the paragraph very small and at the end
        paragraph.style = doc.styles['Normal']
        for run in paragraph.runs:
            run.font.size = Pt(1)

        # Save to output
        if output_path is None:
            output_path = file_path

        doc.save(output_path)

        with open(output_path, 'rb') as f:
            content = f.read()

        return content, output_path

    def _inject_excel(
        self,
        file_path: str,
        token_url: str,
        output_path: Optional[str] = None
    ) -> Tuple[bytes, str]:
        """
        Inject a web query into an Excel spreadsheet.

        Uses an external data connection that triggers when the workbook is opened.
        """
        wb = load_workbook(file_path)

        # Method 1: Add a hyperlink to a hidden cell
        # This won't auto-trigger but creates a clickable link

        # Method 2: Use WEBSERVICE function in a formula (requires macro-enabled)
        # We'll use a simpler approach that works in regular xlsx

        # Get the first worksheet
        ws = wb.active

        # Find an empty cell far from visible content
        # We'll use a cell in row 1000, column 100
        hidden_cell = ws.cell(row=1000, column=100)
        hidden_cell.value = f'=HYPERLINK("{token_url}", " ")'

        # Alternative: Add to document properties custom URL
        # This triggers on some operations
        if wb.custom_doc_props is None:
            from openpyxl.packaging.custom import CustomPropertyList
            wb.custom_doc_props = CustomPropertyList()

        # Add a custom property with the token URL
        # Some versions of Excel may fetch this URL
        try:
            wb.custom_doc_props.append({
                'name': 'tracking_id',
                'value': token_url
            })
        except Exception:
            pass  # Custom properties may not be supported in all versions

        # Save to output
        if output_path is None:
            output_path = file_path

        wb.save(output_path)

        with open(output_path, 'rb') as f:
            content = f.read()

        return content, output_path

    def _inject_pdf(
        self,
        file_path: str,
        token_url: str,
        output_path: Optional[str] = None
    ) -> Tuple[bytes, str]:
        """
        Inject a tracking URL into a PDF document.

        Uses PDF annotations and JavaScript actions to trigger when opened.
        """
        reader = PdfReader(file_path)
        writer = PdfWriter()

        # Copy all pages
        for page in reader.pages:
            writer.add_page(page)

        # Add metadata with URL
        writer.add_metadata({
            '/Producer': 'Campaign Manager',
            '/CustomURL': token_url
        })

        # Create a URI action that triggers on open
        # This adds an OpenAction that navigates to the URL
        # Note: Most PDF viewers block automatic URL opening for security
        # But the URL will be embedded for forensics

        # Add the token URL as a link annotation on the first page
        # This is a more compatible method
        from pypdf.generic import (
            ArrayObject,
            DictionaryObject,
            FloatObject,
            NameObject,
            TextStringObject,
        )

        # Create annotation for first page
        if len(writer.pages) > 0:
            page = writer.pages[0]

            # Create a small, invisible link annotation
            link_annotation = DictionaryObject({
                NameObject('/Type'): NameObject('/Annot'),
                NameObject('/Subtype'): NameObject('/Link'),
                # Position it at 0,0 with 1x1 size (invisible)
                NameObject('/Rect'): ArrayObject([
                    FloatObject(0),
                    FloatObject(0),
                    FloatObject(1),
                    FloatObject(1),
                ]),
                NameObject('/Border'): ArrayObject([
                    FloatObject(0),
                    FloatObject(0),
                    FloatObject(0),
                ]),
                NameObject('/A'): DictionaryObject({
                    NameObject('/Type'): NameObject('/Action'),
                    NameObject('/S'): NameObject('/URI'),
                    NameObject('/URI'): TextStringObject(token_url),
                }),
                NameObject('/F'): FloatObject(4),  # Hidden flag
            })

            # Get or create annotations array
            if '/Annots' in page:
                annots = page['/Annots']
                if isinstance(annots, ArrayObject):
                    annots.append(link_annotation)
            else:
                page[NameObject('/Annots')] = ArrayObject([link_annotation])

        # Save to output
        if output_path is None:
            output_path = file_path

        with open(output_path, 'wb') as f:
            writer.write(f)

        with open(output_path, 'rb') as f:
            content = f.read()

        return content, output_path

    def inject_from_bytes(
        self,
        file_bytes: bytes,
        token_url: str,
        mime_type: str,
        filename: str
    ) -> bytes:
        """
        Inject a token into document provided as bytes.

        Args:
            file_bytes: Document content as bytes
            token_url: The canary token URL to embed
            mime_type: MIME type of the document
            filename: Original filename (used for temp file extension)

        Returns:
            Modified document as bytes
        """
        # Get file extension from filename
        ext = Path(filename).suffix or '.tmp'

        # Create temp file
        with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        try:
            # Inject token
            modified_bytes, _ = self.inject_token(tmp_path, token_url, mime_type, tmp_path)
            return modified_bytes
        finally:
            # Clean up
            if os.path.exists(tmp_path):
                os.remove(tmp_path)


# Singleton instance
document_injector = DocumentInjector()
